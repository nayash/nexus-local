import csv
import hashlib
import json
import mimetypes
import os
import pwd
import re
import uuid
from datetime import datetime
from io import BytesIO
from typing import Callable, Optional

import pandas as pd
from langchain_ollama import OllamaEmbeddings

from src.core.config import Config
from src.embeddings.multimodal_onnx import get_multimodal_embedder
from src.rag.metadata_taxonomy import normalize_document_kind
from src.rag.query_filters import compile_multimodal_filter_plan
from src.rag.lancedb_store import delete_rows, load_rows, search as search_rows, upsert_rows
from src.rag.schemas import build_child_row, build_parent_row, build_registry_row, parse_extra


SUPPORTED_MULTIMODAL_EXTENSIONS = (
    ".pdf",
    ".docx",
    ".txt",
    ".log",
    ".md",
    ".csv",
    ".png",
    ".jpg",
    ".jpeg",
    ".html",
    ".htm",
)

_NOMIC_EMBEDDER = OllamaEmbeddings(
    model="nomic-embed-text",
    base_url=Config.OLLAMA_BASE_URL,
)


def _compute_file_hash(path: str) -> str:
    stat = os.stat(path)
    if stat.st_size <= 20 * 1024 * 1024:
        hasher = hashlib.sha256()
        with open(path, "rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                hasher.update(chunk)
        return hasher.hexdigest()

    composite = f"{os.path.abspath(path)}:{stat.st_mtime_ns}:{stat.st_size}"
    return hashlib.sha256(composite.encode("utf-8")).hexdigest()


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    cleaned = (text or "").strip()
    if not cleaned:
        return []

    if len(cleaned) <= chunk_size:
        return [cleaned]

    chunks = []
    start = 0
    while start < len(cleaned):
        end = min(start + chunk_size, len(cleaned))
        chunk = cleaned[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(cleaned):
            break
        start = max(end - overlap, start + 1)
    return chunks


def _parent_chunks(text: str) -> list[str]:
    return _split_text(text, chunk_size=2000, overlap=200)


def _child_chunks(text: str) -> list[str]:
    return _split_text(text, chunk_size=400, overlap=50)


def _detect_source_type(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    return {
        ".pdf": "pdf",
        ".docx": "docx",
        ".txt": "txt",
        ".log": "log",
        ".md": "md",
        ".csv": "csv",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".html": "html",
        ".htm": "html",
    }.get(ext, ext.lstrip(".") or "unknown")


def _get_file_owner(path: str) -> str:
    try:
        return pwd.getpwuid(os.stat(path).st_uid).pw_name
    except Exception:
        return "unknown"


def _derive_title_author(file_name: str, source_type: str, raw_title: str = "", raw_author: str = "") -> tuple[str, str]:
    title = (raw_title or "").strip()
    author = (raw_author or "").strip()

    stem = os.path.splitext(file_name)[0]
    if not title:
        title = stem

    if not author:
        for separator in (" - ", " — ", " by "):
            if separator in stem:
                left, right = stem.split(separator, 1)
                if separator.strip() == "by":
                    title = left.strip() or title
                    author = right.strip()
                else:
                    author = left.strip()
                    title = right.strip() or title
                break

    if not author and source_type == "docx":
        author = ""
    return title, author


def _guess_document_kind(path: str, source_type: str, title: str, author: str) -> str:
    file_name = os.path.basename(path).lower()
    ext = os.path.splitext(path)[1].lower()
    path_lower = path.lower()
    if source_type == "image":
        kind = "image"
    elif ext == ".sh":
        kind = "code"
    elif ext == ".csv":
        kind = "data"
    elif ext == ".log" or "log" in file_name:
        kind = "log"
    elif "book" in path_lower or "library" in path_lower:
        kind = "book"
    elif author:
        kind = "book"
    elif source_type in {"pdf", "docx", "txt", "md", "html"} and any(token in title.lower() for token in ("chapter", "novel", "poems", "stories")):
        kind = "book"
    else:
        kind = "document"
    return normalize_document_kind(kind, default="document")


def _metadata_with_text_hints(metadata: dict, text: str) -> dict:
    enriched = dict(metadata)
    enriched["document_kind"] = normalize_document_kind(enriched.get("document_kind"), default="document")
    if not text:
        return enriched

    snippet = text[:800]
    if not enriched.get("author"):
        match = re.search(r"\bby\s+([A-Z][A-Za-z]+(?:\s+[A-Z][A-Za-z]+){0,3})", snippet)
        if match:
            enriched["author"] = match.group(1).strip()
            if enriched.get("document_kind") == "document":
                enriched["document_kind"] = normalize_document_kind("book", default="document")
    if enriched.get("title") == os.path.splitext(enriched.get("file_name", ""))[0]:
        first_line = next((line.strip() for line in snippet.splitlines() if line.strip()), "")
        if 3 <= len(first_line) <= 120:
            enriched["title"] = first_line
    return enriched


def _base_metadata(path: str, source_type: str, doc_hash: str, workspace_id: str = "global", raw_title: str = "", raw_author: str = "") -> dict:
    stat = os.stat(path)
    file_name = os.path.basename(path)
    title, author = _derive_title_author(file_name, source_type, raw_title=raw_title, raw_author=raw_author)
    return {
        "workspace_id": workspace_id,
        "file_name": file_name,
        "file_ext": os.path.splitext(file_name)[1].lower(),
        "title": title,
        "author": author,
        "owner": _get_file_owner(path),
        "document_kind": normalize_document_kind(
            _guess_document_kind(path, source_type, title, author),
            default="document",
        ),
        "source_mtime_epoch": int(stat.st_mtime),
        "source_mtime_date": datetime.fromtimestamp(stat.st_mtime).date().isoformat(),
        "source_ctime_epoch": int(stat.st_ctime),
        "source_ctime_date": datetime.fromtimestamp(stat.st_ctime).date().isoformat(),
        "source_size_bytes": int(stat.st_size),
        "doc_hash": doc_hash,
    }


def _extract_file_level_metadata(path: str, source_type: str) -> tuple[str, str]:
    raw_title = ""
    raw_author = ""
    try:
        if source_type == "pdf":
            import fitz

            pdf = fitz.open(path)
            metadata = pdf.metadata or {}
            raw_title = metadata.get("title", "") or ""
            raw_author = metadata.get("author", "") or ""
            pdf.close()
        elif source_type == "docx":
            from docx import Document as DocxDocument

            document = DocxDocument(path)
            props = document.core_properties
            raw_title = getattr(props, "title", "") or ""
            raw_author = getattr(props, "author", "") or ""
    except Exception:
        pass
    return raw_title, raw_author


def _escape_sql(value: str) -> str:
    return value.replace("'", "''")


def _parent_row(
    *,
    parent_id: str,
    doc_hash: str,
    modality: str,
    text: str,
    source_path: str,
    source_type: str,
    page: Optional[int] = None,
    parent_index: Optional[int] = None,
    image_index: Optional[int] = None,
    mime: Optional[str] = None,
    width: Optional[int] = None,
    height: Optional[int] = None,
    extra: Optional[dict] = None,
) -> dict:
    return build_parent_row(
        parent_id=parent_id,
        doc_hash=doc_hash,
        modality=modality,
        text=text,
        source_path=source_path,
        source_type=source_type,
        metadata=dict(extra or {}),
        page=page,
        parent_index=parent_index,
        image_index=image_index,
        mime=mime,
        width=width,
        height=height,
    )


def _child_row(
    *,
    parent_id: str,
    doc_hash: str,
    vector: list[float],
    embedding_family: str,
    modality: str,
    text: str,
    source_path: str,
    source_type: str,
    page: Optional[int] = None,
    parent_index: Optional[int] = None,
    chunk_index: Optional[int] = None,
    image_index: Optional[int] = None,
    mime: Optional[str] = None,
    width: Optional[int] = None,
    height: Optional[int] = None,
    extra: Optional[dict] = None,
) -> dict:
    return build_child_row(
        parent_id=parent_id,
        doc_hash=doc_hash,
        vector=vector,
        embedding_family=embedding_family,
        modality=modality,
        text=text,
        source_path=source_path,
        source_type=source_type,
        metadata=dict(extra or {}),
        page=page,
        parent_index=parent_index,
        chunk_index=chunk_index,
        image_index=image_index,
        mime=mime,
        width=width,
        height=height,
    )


def _save_cached_image(image, doc_hash: str, stem: str) -> str:
    cache_dir = os.path.join(Config.MULTIMODAL_IMAGE_CACHE_DIR, doc_hash)
    os.makedirs(cache_dir, exist_ok=True)
    path = os.path.join(cache_dir, f"{stem}.png")
    image.save(path, format="PNG")
    return path


def _load_registry_row(source_path: str):
    abs_path = os.path.abspath(source_path)
    for row in load_rows(Config.MULTIMODAL_DOCUMENTS_TABLE):
        if row.get("source_path") == abs_path:
            return row
    return None


def _update_registry(source_path: str, source_type: str, doc_hash: str, num_parents: int, num_children: int, metadata: Optional[dict] = None):
    metadata = metadata or {}
    row = build_registry_row(
        source_path=source_path,
        source_type=source_type,
        doc_hash=doc_hash,
        num_parents=num_parents,
        num_children=num_children,
        metadata=metadata,
    )
    escaped = _escape_sql(os.path.abspath(source_path))
    upsert_rows(
        Config.MULTIMODAL_DOCUMENTS_TABLE,
        [row],
        delete_filter=f"source_path = '{escaped}'",
    )


def _append_text_parent(
    parents: list[dict],
    nomic_children: list[dict],
    clip_children: list[dict],
    *,
    source_path: str,
    source_type: str,
    doc_hash: str,
    text: str,
    page: Optional[int],
    parent_index: int,
    extra: Optional[dict] = None,
):
    if not text.strip():
        return

    clip_embedder = get_multimodal_embedder()
    metadata = dict(extra or {})
    for local_parent_index, parent_text in enumerate(_parent_chunks(text)):
        absolute_parent_index = parent_index + local_parent_index
        parent_id = str(uuid.uuid4())
        parent_extra = dict(metadata)
        parents.append(
            _parent_row(
                parent_id=parent_id,
                doc_hash=doc_hash,
                modality="text",
                text=parent_text,
                source_path=source_path,
                source_type=source_type,
                page=page,
                parent_index=absolute_parent_index,
                extra=parent_extra,
            )
        )

        child_texts = _child_chunks(parent_text)
        if child_texts:
            nomic_vectors = _NOMIC_EMBEDDER.embed_documents(child_texts)
            for child_index, (child_text, vector) in enumerate(zip(child_texts, nomic_vectors)):
                nomic_children.append(
                    _child_row(
                        parent_id=parent_id,
                        doc_hash=doc_hash,
                        vector=vector,
                        embedding_family="nomic",
                        modality="text",
                        text=child_text,
                        source_path=source_path,
                        source_type=source_type,
                        page=page,
                        parent_index=absolute_parent_index,
                        chunk_index=child_index,
                        extra=parent_extra,
                    )
                )

            if clip_embedder:
                clip_vectors = clip_embedder.embed_texts(child_texts)
                for child_index, (child_text, vector) in enumerate(zip(child_texts, clip_vectors)):
                    clip_children.append(
                        _child_row(
                            parent_id=parent_id,
                            doc_hash=doc_hash,
                            vector=vector,
                            embedding_family="clip",
                            modality="text",
                            text=child_text,
                            source_path=source_path,
                            source_type=source_type,
                            page=page,
                            parent_index=absolute_parent_index,
                            chunk_index=child_index,
                            extra=parent_extra,
                        )
                    )


def _append_image_parent(
    parents: list[dict],
    clip_children: list[dict],
    *,
    source_path: str,
    source_type: str,
    doc_hash: str,
    image,
    page: Optional[int],
    image_index: int,
    mime: str,
    extra: Optional[dict] = None,
):
    clip_embedder = get_multimodal_embedder()
    if not clip_embedder:
        return

    parent_id = str(uuid.uuid4())
    parent_extra = dict(extra or {})
    parents.append(
        _parent_row(
            parent_id=parent_id,
            doc_hash=doc_hash,
            modality="image",
            text="",
            source_path=source_path,
            source_type=source_type,
            page=page,
            image_index=image_index,
            mime=mime,
            width=image.width,
            height=image.height,
            extra=parent_extra,
        )
    )
    clip_children.append(
        _child_row(
            parent_id=parent_id,
            doc_hash=doc_hash,
            vector=clip_embedder.embed_image(image),
            embedding_family="clip",
            modality="image",
            text=parent_extra.get("caption", ""),
            source_path=source_path,
            source_type=source_type,
            page=page,
            image_index=image_index,
            mime=mime,
            width=image.width,
            height=image.height,
            extra=parent_extra,
        )
    )


def _extract_pdf_content(path: str, doc_hash: str, base_metadata: dict) -> tuple[list[dict], list[dict], list[dict]]:
    parents: list[dict] = []
    nomic_children: list[dict] = []
    clip_children: list[dict] = []
    try:
        import fitz
        from PIL import Image
    except ImportError:
        from langchain_community.document_loaders import PyPDFLoader

        loader = PyPDFLoader(path)
        docs = loader.load()
        parent_index = 0
        for page_number, doc in enumerate(docs, start=1):
            page_metadata = _metadata_with_text_hints(base_metadata, doc.page_content)
            _append_text_parent(
                parents,
                nomic_children,
                clip_children,
                source_path=path,
                source_type="pdf",
                doc_hash=doc_hash,
                text=doc.page_content,
                page=page_number,
                parent_index=parent_index,
                extra={**page_metadata, "fallback": "pypdf"},
            )
            parent_index = len([row for row in parents if row.get("modality") == "text"])
        return parents, nomic_children, clip_children

    pdf = fitz.open(path)
    text_parent_index = 0
    for page_index in range(len(pdf)):
        page = pdf.load_page(page_index)
        page_text = page.get_text("text")
        page_metadata = _metadata_with_text_hints(base_metadata, page_text)
        _append_text_parent(
            parents,
            nomic_children,
            clip_children,
            source_path=path,
            source_type="pdf",
            doc_hash=doc_hash,
            text=page_text,
            page=page_index + 1,
            parent_index=text_parent_index,
            extra=page_metadata,
        )
        text_parent_index = len([row for row in parents if row.get("modality") == "text"])

        for image_idx, image_meta in enumerate(page.get_images(full=True)):
            extracted = pdf.extract_image(image_meta[0])
            image_bytes = extracted.get("image")
            if not image_bytes:
                continue
            image = Image.open(BytesIO(image_bytes)).convert("RGB")
            cached_path = _save_cached_image(image, doc_hash, f"p{page_index + 1}_{image_idx}")
            _append_image_parent(
                parents,
                clip_children,
                source_path=path,
                source_type="pdf",
                doc_hash=doc_hash,
                image=image,
                page=page_index + 1,
                image_index=image_idx,
                mime="image/png",
            extra={**page_metadata, "cached_path": cached_path},
            )
    pdf.close()
    return parents, nomic_children, clip_children


def _extract_docx_content(path: str, doc_hash: str, base_metadata: dict) -> tuple[list[dict], list[dict], list[dict]]:
    parents: list[dict] = []
    nomic_children: list[dict] = []
    clip_children: list[dict] = []
    try:
        from docx import Document as DocxDocument
        from PIL import Image
    except ImportError:
        return parents, nomic_children, clip_children

    document = DocxDocument(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    text_metadata = _metadata_with_text_hints(base_metadata, text)
    _append_text_parent(
        parents,
        nomic_children,
        clip_children,
        source_path=path,
        source_type="docx",
        doc_hash=doc_hash,
        text=text,
        page=None,
        parent_index=0,
        extra=text_metadata,
    )

    image_index = 0
    for rel in document.part.rels.values():
        target_part = getattr(rel, "target_part", None)
        blob = getattr(target_part, "blob", None)
        if not blob:
            continue
        content_type = getattr(target_part, "content_type", "")
        if not content_type.startswith("image/"):
            continue
        image = Image.open(BytesIO(blob)).convert("RGB")
        cached_path = _save_cached_image(image, doc_hash, f"docx_{image_index}")
        _append_image_parent(
            parents,
            clip_children,
            source_path=path,
            source_type="docx",
            doc_hash=doc_hash,
            image=image,
            page=None,
            image_index=image_index,
            mime="image/png",
            extra={**text_metadata, "cached_path": cached_path},
        )
        image_index += 1
    return parents, nomic_children, clip_children


def _read_text_file(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8") as handle:
            return handle.read()
    except UnicodeDecodeError:
        with open(path, "r", encoding="utf-8", errors="ignore") as handle:
            return handle.read()


def _extract_text_content(path: str, source_type: str, doc_hash: str, base_metadata: dict) -> tuple[list[dict], list[dict], list[dict]]:
    text = _read_text_file(path)
    if source_type == "html":
        try:
            from bs4 import BeautifulSoup

            text = BeautifulSoup(text, "html.parser").get_text("\n")
        except ImportError:
            text = re.sub(r"<[^>]+>", " ", text)
    text_metadata = _metadata_with_text_hints(base_metadata, text)

    parents: list[dict] = []
    nomic_children: list[dict] = []
    clip_children: list[dict] = []
    _append_text_parent(
        parents,
        nomic_children,
        clip_children,
        source_path=path,
        source_type=source_type,
        doc_hash=doc_hash,
        text=text,
        page=None,
        parent_index=0,
        extra=text_metadata,
    )
    return parents, nomic_children, clip_children


def _extract_csv_content(path: str, doc_hash: str, base_metadata: dict, rows_per_parent: int = 200) -> tuple[list[dict], list[dict], list[dict]]:
    parents: list[dict] = []
    nomic_children: list[dict] = []
    clip_children: list[dict] = []
    try:
        frame = pd.read_csv(path)
    except Exception:
        with open(path, "r", encoding="utf-8", errors="ignore", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
        if not rows:
            return parents, nomic_children, clip_children
        frame = pd.DataFrame(rows[1:], columns=rows[0])

    for parent_index, start in enumerate(range(0, len(frame), rows_per_parent)):
        end = min(start + rows_per_parent, len(frame))
        parent_text = frame.iloc[start:end].to_csv(index=False)
        row_metadata = _metadata_with_text_hints(base_metadata, parent_text)
        _append_text_parent(
            parents,
            nomic_children,
            clip_children,
            source_path=path,
            source_type="csv",
            doc_hash=doc_hash,
            text=parent_text,
            page=None,
            parent_index=parent_index,
            extra={**row_metadata, "row_start": int(start), "row_end": int(end - 1)},
        )
    return parents, nomic_children, clip_children


def _extract_image_content(path: str, doc_hash: str, base_metadata: dict) -> tuple[list[dict], list[dict], list[dict]]:
    parents: list[dict] = []
    nomic_children: list[dict] = []
    clip_children: list[dict] = []
    try:
        from PIL import Image
    except ImportError:
        return parents, nomic_children, clip_children

    image = Image.open(path).convert("RGB")
    mime = mimetypes.guess_type(path)[0] or "image/png"
    cached_path = _save_cached_image(image, doc_hash, "image_0")
    _append_image_parent(
        parents,
        clip_children,
        source_path=path,
        source_type="image",
        doc_hash=doc_hash,
        image=image,
        page=None,
        image_index=0,
        mime=mime,
        extra={**base_metadata, "cached_path": cached_path},
    )
    return parents, nomic_children, clip_children


def _extract_content(path: str, source_type: str, doc_hash: str, base_metadata: dict) -> tuple[list[dict], list[dict], list[dict]]:
    if source_type == "pdf":
        return _extract_pdf_content(path, doc_hash, base_metadata)
    if source_type == "docx":
        return _extract_docx_content(path, doc_hash, base_metadata)
    if source_type in {"txt", "log", "md", "html"}:
        return _extract_text_content(path, source_type, doc_hash, base_metadata)
    if source_type == "csv":
        return _extract_csv_content(path, doc_hash, base_metadata)
    if source_type == "image":
        return _extract_image_content(path, doc_hash, base_metadata)
    return [], [], []


def ingest_file_multimodal(path: str, progress_callback: Optional[Callable] = None) -> tuple[bool, int, Optional[str]]:
    abs_path = os.path.abspath(os.path.expanduser(path))
    if not os.path.isfile(abs_path):
        return False, 0, None

    if os.path.splitext(abs_path)[1].lower() not in SUPPORTED_MULTIMODAL_EXTENSIONS:
        return False, 0, None

    source_type = _detect_source_type(abs_path)
    doc_hash = _compute_file_hash(abs_path)
    raw_title, raw_author = _extract_file_level_metadata(abs_path, source_type)
    base_metadata = _base_metadata(abs_path, source_type, doc_hash, raw_title=raw_title, raw_author=raw_author)
    existing = _load_registry_row(abs_path)
    if existing and existing.get("hash") == doc_hash:
        count = int(existing.get("num_children") or existing.get("num_chunks") or 0)
        return True, count, doc_hash

    parents, nomic_children, clip_children = _extract_content(abs_path, source_type, doc_hash, base_metadata)
    total_children = len(nomic_children) + len(clip_children)

    escaped_path = _escape_sql(abs_path)
    delete_filter = f"source_path = '{escaped_path}'"
    delete_rows(Config.MULTIMODAL_PARENT_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_TEXT_CHILD_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_CLIP_CHILD_TABLE, delete_filter)

    if parents:
        upsert_rows(
            Config.MULTIMODAL_PARENT_TABLE,
            parents,
            delete_filter=None,
        )

    if nomic_children:
        upsert_rows(
            Config.MULTIMODAL_TEXT_CHILD_TABLE,
            nomic_children,
            delete_filter=None,
        )

    if clip_children:
        upsert_rows(
            Config.MULTIMODAL_CLIP_CHILD_TABLE,
            clip_children,
            delete_filter=None,
        )

    _update_registry(abs_path, source_type, doc_hash, len(parents), total_children, metadata=base_metadata)

    if progress_callback:
        progress_callback()

    return True, total_children, doc_hash


def ingest_paths(paths: list[str], db_path: str = "", table_name: str = "", embedder=None):
    _ = db_path, table_name, embedder
    successes = 0
    total_rows = 0
    for path in paths:
        success, row_count, _ = ingest_file_multimodal(path)
        if success:
            successes += 1
            total_rows += row_count
    return successes, total_rows


def purge_multimodal_rows(source_path: str):
    abs_path = os.path.abspath(source_path)
    escaped = _escape_sql(abs_path)
    delete_filter = f"source_path = '{escaped}'"
    delete_rows(Config.MULTIMODAL_PARENT_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_TEXT_CHILD_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_CLIP_CHILD_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_DOCUMENTS_TABLE, delete_filter)


def purge_multimodal_prefix(path_prefix: str):
    abs_prefix = os.path.abspath(path_prefix)
    escaped = _escape_sql(abs_prefix)
    delete_filter = f"source_path LIKE '{escaped}%'"
    delete_rows(Config.MULTIMODAL_PARENT_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_TEXT_CHILD_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_CLIP_CHILD_TABLE, delete_filter)
    delete_rows(Config.MULTIMODAL_DOCUMENTS_TABLE, delete_filter)


def is_source_indexed_multimodal(source_path: str) -> bool:
    abs_path = os.path.abspath(source_path)
    for row in load_rows(Config.MULTIMODAL_DOCUMENTS_TABLE):
        if row.get("source_path") == abs_path:
            return True
    return False


def _query_prefers_images(query: str) -> bool:
    lowered = query.lower()
    keywords = ("only images", "image only", "diagram", "screenshot", "photo", "picture", "chart", "graph")
    return any(token in lowered for token in keywords)


def _query_is_text_heavy(query: str) -> bool:
    lowered = query.lower()
    image_only_terms = ("show me the image", "show me the screenshot", "find the image")
    return not any(term in lowered for term in image_only_terms)


def _search_nomic_children(query: str, top_k: int, sql_filter: Optional[str]) -> list[dict]:
    vector = _NOMIC_EMBEDDER.embed_query(query)
    rows = search_rows(
        Config.MULTIMODAL_TEXT_CHILD_TABLE,
        vector,
        filters=sql_filter,
        top_k=top_k,
    )
    for index, row in enumerate(rows):
        row["_retrieval_channel"] = "nomic"
        row["_rank"] = index
    return rows


def _search_clip_children(query: str, top_k: int, sql_filter: Optional[str]) -> list[dict]:
    clip_embedder = get_multimodal_embedder()
    if not clip_embedder:
        return []

    rows = search_rows(
        Config.MULTIMODAL_CLIP_CHILD_TABLE,
        clip_embedder.embed_text(query),
        filters=sql_filter,
        top_k=top_k,
    )
    for index, row in enumerate(rows):
        row["_retrieval_channel"] = "clip"
        row["_rank"] = index
    return rows


_LEXICAL_STOPWORDS = {
    "the",
    "a",
    "an",
    "for",
    "from",
    "with",
    "about",
    "into",
    "in",
    "on",
    "at",
    "to",
    "of",
    "my",
    "me",
    "please",
    "give",
    "find",
    "search",
}


def _query_terms(query: str) -> list[str]:
    terms = []
    for token in re.findall(r"[A-Za-z0-9][A-Za-z0-9_\-]{1,}", query.lower()):
        if token in _LEXICAL_STOPWORDS:
            continue
        if len(token) < 3:
            continue
        terms.append(token)
    return terms


def _build_source_path_in_filter(paths: list[str]) -> Optional[str]:
    if not paths:
        return None
    escaped_paths = [f"'{_escape_sql(os.path.abspath(path))}'" for path in paths]
    if len(escaped_paths) == 1:
        return f"source_path = {escaped_paths[0]}"
    return f"source_path IN ({', '.join(escaped_paths)})"


def _lexical_source_path_filter(query: str, file_filter: Optional[str], limit: int = 20) -> Optional[str]:
    abs_filter = os.path.abspath(file_filter) if file_filter else None
    terms = _query_terms(query)
    if not terms:
        return None

    scored = []
    for row in load_rows(Config.MULTIMODAL_DOCUMENTS_TABLE):
        source_path = os.path.abspath(row.get("source_path") or "")
        if not source_path:
            continue
        if abs_filter and source_path != abs_filter:
            continue

        haystack = " ".join(
            [
                (row.get("file_name") or ""),
                (row.get("title") or ""),
                source_path,
            ]
        ).lower()
        score = sum(1 for term in terms if term in haystack)
        if score <= 0:
            continue
        scored.append((score, source_path))

    if not scored:
        return None

    ranked_paths = [path for _, path in sorted(scored, key=lambda item: (-item[0], item[1]))]
    return _build_source_path_in_filter(ranked_paths[:limit])


def _row_score(row: dict, query_prefers_images: bool) -> float:
    distance = row.get("_distance")
    if distance is None:
        base = 1.0 / (1 + row.get("_rank", 0))
    else:
        try:
            base = 1.0 / (1.0 + float(distance))
        except Exception:
            base = 1.0 / (1 + row.get("_rank", 0))

    channel = row.get("_retrieval_channel")
    modality = row.get("modality")
    if channel == "nomic" and modality == "text":
        base += 0.15
    if channel == "clip" and modality == "image":
        base += 0.1
    if query_prefers_images and modality == "image":
        base += 0.2
    if query_prefers_images and channel == "clip":
        base += 0.05
    return base


def _load_parent_lookup(file_filter: Optional[str]) -> dict[str, dict]:
    parents = load_rows(Config.MULTIMODAL_PARENT_TABLE)
    lookup = {}
    abs_filter = os.path.abspath(file_filter) if file_filter else None
    for row in parents:
        source_path = row.get("source_path")
        if abs_filter and os.path.abspath(source_path or "") != abs_filter:
            continue
        lookup[row.get("parent_id")] = row
    return lookup


def _run_semantic_retrieval_pass(
    *,
    label: str,
    semantic_query: str,
    sql_filter: Optional[str],
    nomic_pool: int,
    clip_pool: int,
) -> list[dict]:
    child_hits = []
    nomic_hits = 0
    clip_hits = 0

    try:
        nomic_rows = _search_nomic_children(semantic_query, top_k=nomic_pool, sql_filter=sql_filter)
        nomic_hits = len(nomic_rows)
        child_hits.extend(nomic_rows)
    except Exception as exc:
        print(f"⚠️ Nomic child search failed ({label} pass): {exc}")

    try:
        clip_rows = _search_clip_children(semantic_query, top_k=clip_pool, sql_filter=sql_filter)
        clip_hits = len(clip_rows)
        child_hits.extend(clip_rows)
    except Exception as exc:
        print(f"⚠️ CLIP child search failed ({label} pass): {exc}")

    print(
        "local retrieval pass | "
        f"label={label} | sql_filter={sql_filter!r} | "
        f"nomic_hits={nomic_hits} | clip_hits={clip_hits} | total_hits={len(child_hits)}"
    )
    return child_hits


def search_multimodal(query: str, top_k: int = 5, file_filter: Optional[str] = None, workspace_id: Optional[str] = None) -> list[dict]:
    query_prefers_images = _query_prefers_images(query)
    query_text_heavy = _query_is_text_heavy(query)
    plan = compile_multimodal_filter_plan(query, file_filter=file_filter, workspace_id=workspace_id)
    semantic_query = plan.text_query

    nomic_pool = max(top_k * 3, top_k) if query_text_heavy else max(top_k * 2, top_k)
    clip_pool = max(top_k * 3, top_k) if query_prefers_images else max(top_k * 2, top_k)

    if plan.strict_clauses:
        clause_summary = [
            f"{item.attribute} {item.comparator} {item.value} ({item.confidence})"
            for item in plan.strict_clauses
        ]
        print(f"local retrieval filter clauses: {clause_summary}")
    if plan.dropped_clauses:
        dropped_summary = [f"{item.attribute} {item.comparator} {item.value}" for item in plan.dropped_clauses]
        print(f"local retrieval low-confidence clauses: {dropped_summary}")

    attempts: list[tuple[str, Optional[str]]] = [("strict", plan.strict_sql_filter)]
    if plan.should_try_relaxed:
        attempts.append(("relaxed", plan.relaxed_sql_filter))
    if attempts[-1][1] is not None:
        attempts.append(("unfiltered", None))

    child_hits = []
    for label, sql_filter in attempts:
        child_hits = _run_semantic_retrieval_pass(
            label=label,
            semantic_query=semantic_query,
            sql_filter=sql_filter,
            nomic_pool=nomic_pool,
            clip_pool=clip_pool,
        )
        if child_hits:
            break

    if not child_hits:
        lexical_filter = _lexical_source_path_filter(query, file_filter=file_filter)
        if lexical_filter and lexical_filter not in {sql for _, sql in attempts}:
            child_hits = _run_semantic_retrieval_pass(
                label="lexical+semantic",
                semantic_query=semantic_query,
                sql_filter=lexical_filter,
                nomic_pool=nomic_pool,
                clip_pool=clip_pool,
            )

    if not child_hits:
        return []

    parent_lookup = _load_parent_lookup(file_filter)
    best_by_parent = {}
    for child in child_hits:
        parent_id = child.get("parent_id")
        parent = parent_lookup.get(parent_id)
        if not parent:
            continue

        score = _row_score(child, query_prefers_images=query_prefers_images)
        existing = best_by_parent.get(parent_id)
        if not existing or score > existing["score"]:
            best_by_parent[parent_id] = {
                "score": score,
                "parent": dict(parent),
                "child": child,
            }

    if not best_by_parent:
        return []

    ranked = sorted(best_by_parent.values(), key=lambda item: item["score"], reverse=True)[:top_k]
    results = []
    for item in ranked:
        parent = item["parent"]
        child = item["child"]
        extra = parse_extra(parent.get("extra"))
        extra["matched_embedding_family"] = child.get("embedding_family")
        extra["matched_modality"] = child.get("modality")
        extra["matched_text"] = child.get("text", "")
        parent["extra"] = json.dumps(extra, ensure_ascii=True)
        parent["_score"] = item["score"]
        results.append(parent)
    return results
